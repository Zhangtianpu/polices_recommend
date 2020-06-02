#!/usr/bin/env python
# coding: utf-8

# In[2]:


from pdfminer.pdfparser import PDFParser, PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.pdfdevice import PDFDevice
from pdfminer.layout import LAParams
from pdfminer.converter import PDFPageAggregator
import jieba.posseg as pseg
import codecs
import docx
import os


# In[4]:


#得到政策数据集
def get_train_file_name(path,fileList):
    fileNames=os.listdir(path)
    for fileName in fileNames:
        newFileName=os.path.join(path,fileName)
        if os.path.isdir(newFileName):
            get_train_file_name(newFileName,fileList)
        else:
            if fileName=='.DS_Store':
                continue
            fileList.append(newFileName)
            
#2构建停用词表 './data/stop_words.txt'
def create_stop_word(stop_word_file_name):
    #datapath = 'E:/Program Files/python/data/testres0519.txt'
    #querypath = 'E:/Program Files/python/data/queryres0519.txt'
    #storepath = 'E:/Program Files/python/data/store0519.txt'
    stop_words =stop_word_file_name
    stopwords = codecs.open(stop_words,'r',encoding='utf8').readlines()
    stopwords = [ w.strip() for w in stopwords ]
    #3shez结巴分词后的停用词性 [标点符号、连词、助词、副词、介词、时语素、‘的’、数词、方位词、代词]
    stop_flag = ['x', 'c', 'u','d', 'p', 't', 'uj', 'm', 'f', 'r']
    return stopwords,stop_flag

def _read_pdf(fp):
    re=""
    parser = PDFParser(fp)
    # PDF 文档的对象
    doc = PDFDocument()
    # 连接解释器与文档对象
    parser.set_document(doc)
    doc.set_parser(parser)
    # 初始化文档
    doc.initialize("")
    # 创建PDF资源管理器
    resource = PDFResourceManager()
    # 参数分析器
    laparam = LAParams()
    # 创建一个聚合器
    device = PDFPageAggregator(resource, laparams = laparam)
    # 页面解释器
    interpreter = PDFPageInterpreter(resource, device)
    # 使用文档对象得到页面的集合
    for page in doc.get_pages():
        # 使用页面解释器来读取
        interpreter.process_page(page)
        # 使用聚合器获得内容
        layout = device.get_result()
        for out in layout:
            if hasattr(out, "get_text"):
                #print(out.get_text())
                re+=out.get_text()
    return re

def read_docx(fileName):
    result=""
    try:
        file=docx.Document(fileName)
        #输出每一段的内容
    except :
        #print(fileName)
        return result
    else:
        for para in file.paragraphs:
            result+=para.text
        return result
    
    
def _tokenization(filename,stopwords,stop_flag):
    result = []
    sub_re=""
    _fileName=filename.split('\\')[-1]
    _postfix=_fileName.split('.')[-1]
    if _postfix=='docx':
        sub_re=read_docx(filename)
    if _postfix=='pdf':
        print(filename)
        with open(filename, "rb") as my_pdf:
            sub_re=_read_pdf(my_pdf)
    words = pseg.cut(sub_re)
    for word, flag in words:
        if flag not in stop_flag and word not in stopwords:
            result.append(word)
    return result

def tokenization_for_search(keywords,stopwords,stop_flag):
    result=[]
    words = pseg.cut(keywords)
    for word, flag in words:
        if flag not in stop_flag and word not in stopwords:
            result.append(word)
    return result

